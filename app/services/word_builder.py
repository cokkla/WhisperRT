"""
Word文档构建器
用于将解析后的Markdown元素转换为Word文档
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from typing import List, Optional
from app.services.markdown_parser import MarkdownElement, ElementType, TextFormat
from docx.oxml.ns import qn

class WordDocumentBuilder:
    """Word文档构建器"""
    
    def __init__(self):
        self.doc = Document()
        self._setup_default_styles()
    
    def _setup_default_styles(self):
        """设置默认样式"""
        # 设置默认字体
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'     # 英文字体
        font.size = Pt(12)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')    # 中文


    def build_document(self, elements: List[MarkdownElement]) -> Document:
        """
        构建Word文档
        
        Args:
            elements: 解析后的Markdown元素列表
            
        Returns:
            构建好的Word文档
        """
        for element in elements:
            self._add_element(element)
        
        return self.doc
    
    def _add_element(self, element: MarkdownElement):
        """添加单个元素到文档"""
        if element.type == ElementType.HEADING:
            self._add_heading(element)
        elif element.type == ElementType.PARAGRAPH:
            self._add_paragraph(element)
        elif element.type == ElementType.TABLE:
            self._add_table(element)
        elif element.type == ElementType.CODE_BLOCK:
            self._add_code_block(element)
        elif element.type == ElementType.HORIZONTAL_RULE:
            self._add_horizontal_rule()
        elif element.type == ElementType.EMPTY_LINE:
            self._add_empty_line()
    
    def _add_heading(self, element: MarkdownElement):
        """添加标题"""
        heading = self.doc.add_heading(level=element.level)

        # 清除默认文本
        heading.clear()

        # 设置标题字体样式
        run = heading.add_run(element.content)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16 - element.level)
        run.bold = True
        run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 认为标题不存在其他样式，如斜体...暂不做其他样式处理
        # if element.formatted_text:
        #     self._add_formatted_text_to_paragraph(heading, element.formatted_text)
        # else:
        #     heading.text = element.content
    
    def _add_paragraph(self, element: MarkdownElement):
        """添加段落"""
        # 检查是否为列表项（通过内容特征判断）
        is_list_item = (element.content.startswith('**') and '**:' in element.content) or \
                      (element.content.startswith('- ') or element.content.startswith('* ') or element.content.startswith('+ '))
        
        if is_list_item:
            # 列表项使用项目符号样式
            paragraph = self.doc.add_paragraph(style='List Bullet')
        else:
            # 普通段落
            paragraph = self.doc.add_paragraph()
        
        if element.formatted_text:
            self._add_formatted_text_to_paragraph(paragraph, element.formatted_text)
        else:
            paragraph.text = element.content
    
    def _add_table(self, element: MarkdownElement):
        """添加表格"""
        if not element.table_data or len(element.table_data) < 2:
            return
        
        # 创建表格
        table = self.doc.add_table(rows=len(element.table_data), cols=len(element.table_data[0]))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 填充表格数据
        for row_idx, row_data in enumerate(element.table_data):
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(table.rows[row_idx].cells):
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = cell_data
                    
                    # 设置表头样式
                    if row_idx == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
        
        # 设置表格列宽
        for col in table.columns:
            col.width = Inches(1.5)
    
    
    def _add_code_block(self, element: MarkdownElement):
        """添加代码块"""
        paragraph = self.doc.add_paragraph()
        paragraph.style = 'No Spacing'
        
        # 设置代码块样式
        run = paragraph.add_run(element.content)
        run.font.name = 'Consolas'  # 使用等宽字体
        run.font.size = Pt(10)
        
        # 设置段落背景色（浅灰色）
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.right_indent = Inches(0.5)
    
    def _add_horizontal_rule(self):
        """添加水平线"""
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加一行连字符作为水平线
        run = paragraph.add_run('─' * 50)
        run.font.size = Pt(8)
    
    def _add_empty_line(self):
        """添加空行"""
        self.doc.add_paragraph()
    
    def _add_formatted_text_to_paragraph(self, paragraph, formatted_text: List[TextFormat]):
        """向段落添加格式化文本"""
        for text_format in formatted_text:
            run = paragraph.add_run(text_format.text)
            
            if text_format.bold:
                run.bold = True
            if text_format.italic:
                run.italic = True
    
    def save_document(self, file_path: str):
        """保存文档"""
        self.doc.save(file_path)
    
    def get_document(self) -> Document:
        """获取文档对象"""
        return self.doc
