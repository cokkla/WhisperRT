"""
会议纪要生成API端点
"""
import asyncio
import json
import os.path
import uuid
import logging
from typing import Dict, Any, Optional
from datetime import datetime
from fastapi import APIRouter, HTTPException, BackgroundTasks
from fastapi.responses import StreamingResponse, FileResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Inches
from app.core.logging import logger
from app.config import FILE_DOWNLOAD_CONFIG
from app.services.meeting_assistant import MeetingAssistantAgent
from app.services.llm import ChatLLM
from app.models.schemas import MeetingInfo, MeetingMinutesResponse, MeetingMinutesRequest

router = APIRouter()

# 全局存储任务状态和结果
task_storage: Dict[str, Dict[str, Any]] = {}

@router.post("/generate", response_model=MeetingMinutesResponse)
async def generate_meeting_minutes(request: MeetingMinutesRequest):
    """
    生成会议纪要
    
    Args:
        request: 会议纪要生成请求
        
    Returns:
        任务ID和状态信息
    """
    try:
        # 生成任务ID
        task_id = str(uuid.uuid4())
        
        # 初始化会议助手
        chat_llm = ChatLLM()
        meeting_assistant = MeetingAssistantAgent(chat_llm)
        
        # 准备会议信息
        meeting_info = {
            "meeting_topic": request.meeting_info.topic,
            "meeting_date": request.meeting_info.date,
            "meeting_location": request.meeting_info.location,
            "attendees": ", ".join(request.meeting_info.attendees),
            "meeting_host": request.meeting_info.host,
            "recorder": request.meeting_info.recorder
        }
        
        # 存储任务信息
        task_storage[task_id] = {
            "status": "processing",
            "meeting_info": meeting_info,
            "transcription_text": request.transcription_text,
            "created_at": datetime.now().isoformat(),
            "result": None
        }
        
        return MeetingMinutesResponse(
            task_id=task_id,
            status="processing",
            message="会议纪要生成任务已创建"
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"生成会议纪要失败: {str(e)}")


@router.get("/stream/{task_id}")
async def stream_meeting_minutes(task_id: str):
    """
    流式获取会议纪要生成进度
    
    Args:
        task_id: 任务ID
        
    Returns:
        Server-Sent Events流
    """
    if task_id not in task_storage:
        raise HTTPException(status_code=404, detail="任务不存在")
    
    task_info = task_storage[task_id]
    
    async def generate_stream():
        try:
            # 初始化会议助手
            chat_llm = ChatLLM()
            meeting_assistant = MeetingAssistantAgent(chat_llm)
            
            # 更新任务状态
            task_storage[task_id]["status"] = "processing"
            
            # 流式处理会议内容
            async for result in meeting_assistant.process_meeting_stream(
                task_info["transcription_text"],
                task_info["meeting_info"]
            ):
                # 发送SSE数据
                yield f"data: {json.dumps(result, ensure_ascii=False)}\n\n"
                await asyncio.sleep(0.01)

                # 如果是最终结果，保存到任务存储
                if result.get("step") == "generating" and result.get("status") == "completed":
                    task_storage[task_id]["status"] = "completed"
                    # 保存完整的会议纪要内容
                    content = result.get("content", "")
                    task_storage[task_id]["result"] = content["会议纪要"]
            yield "event: done\ndata: 会议纪要生成完成\n\n"

        except Exception as e:
            # 发送错误信息
            error_result = {
                "step": "error",
                "status": "failed",
                "message": f"处理失败: {str(e)}"
            }
            yield f"data: {json.dumps(error_result, ensure_ascii=False)}\n\n"
            task_storage[task_id]["status"] = "failed"
            task_storage[task_id]["error"] = str(e)

        # 心跳机制 防止前端长时间未接收到数据，部分浏览器或代理会自动断开连接
        while task_storage[task_id]["status"] == "processing":
            yield "data: {}\n\n"
            await asyncio.sleep(15)
    
    return StreamingResponse(
        generate_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
        }
    )


@router.get("/status/{task_id}")
async def get_task_status(task_id: str):
    """
    获取任务状态
    
    Args:
        task_id: 任务ID
        
    Returns:
        任务状态信息
    """
    if task_id not in task_storage:
        raise HTTPException(status_code=404, detail="任务不存在")
    
    task_info = task_storage[task_id]
    return {
        "task_id": task_id,
        "status": task_info["status"],
        "created_at": task_info["created_at"],
        "result_length": len(task_info.get("result", "")) if task_info.get("result") else 0
    }


@router.get("/download/{task_id}")
async def download_meeting_minutes(task_id: str, format: str = "markdown"):
    """
    下载会议纪要文档
    
    Args:
        task_id: 任务ID
        format: 文档格式 (markdown 或 word)
        
    Returns:
        文档文件
    """
    if task_id not in task_storage:
        raise HTTPException(status_code=404, detail="任务不存在")
    
    task_info = task_storage[task_id]
    
    if task_info["status"] != "completed":
        raise HTTPException(status_code=400, detail="任务尚未完成")
    
    if not task_info.get("result"):
        raise HTTPException(status_code=400, detail="没有可下载的内容")
    
    content = task_info["result"]

    path_root  = FILE_DOWNLOAD_CONFIG["temp_dir"]
    if not os.path.exists(path_root):
        os.makedirs(path_root)
        logger.info(f"创建文件目录: {path_root}")

    if format.lower() == FILE_DOWNLOAD_CONFIG["allowed_output_format"][0]:
        # 生成Markdown文件
        filename = f"meeting_minutes_{task_id[:4]}.md"
        file_path = f"{path_root}/{filename}"
        
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)
        
        return FileResponse(
            path=file_path,
            filename=filename,
            media_type="text/markdown"
        )
    
    elif format.lower() == FILE_DOWNLOAD_CONFIG["allowed_output_format"][1]:
        # 生成Word文件
        doc = Document()

        # 解析Markdown内容并转换为Word格式
        lines = content.split('\n')
        current_paragraph = None

        for line in lines:
            line = line.strip()
            if not line:
                continue

            if line.startswith('# '):
                # 主标题
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                # 二级标题
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                # 三级标题
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- **') and '**:' in line:
                # 列表项
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('- '):
                # 普通列表项
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                # 普通段落
                doc.add_paragraph(line)

        filename = f"meeting_minutes_{task_id[:8]}.docx"
        file_path = f"temp_audio/{filename}"
        doc.save(file_path)

        return FileResponse(
            path=file_path,
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    else:
        raise HTTPException(status_code=400, detail="不支持的格式，请使用markdown或word")


@router.delete("/task/{task_id}")
async def delete_task(task_id: str):
    """
    删除任务
    
    Args:
        task_id: 任务ID
        
    Returns:
        删除结果
    """
    if task_id not in task_storage:
        raise HTTPException(status_code=404, detail="任务不存在")
    
    del task_storage[task_id]
    return {"message": "任务已删除"}


@router.get("/tasks")
async def list_tasks():
    """
    列出所有任务
    
    Returns:
        任务列表
    """
    tasks = []
    for task_id, task_info in task_storage.items():
        tasks.append({
            "task_id": task_id,
            "status": task_info["status"],
            "created_at": task_info["created_at"],
            "result_length": len(task_info.get("result", "")) if task_info.get("result") else 0
        })
    
    return {"tasks": tasks}
